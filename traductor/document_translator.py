"""
Document Translation Agent Module.

Provides the DocumentTranslator class for translating documents
(PDF, DOCX, TXT) using Ollama's translation models.
"""

import os
import logging
from typing import Optional, Dict, List
from pathlib import Path

# Import custom exceptions
from traductor.exceptions import (
    TranslationException,
    TranslationUnsupportedFormatError,
    TranslationAPIError
)

# Import ollama_manager for API calls (located in transcribe directory)
import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'transcribe'))
from transcribe.ollama_manager import call_ollama_api_request, procesa_output, OllamaException


# Supported file formats
SUPPORTED_FORMATS = ['pdf', 'docx', 'txt']

# Language name to code mapping
LANGUAGE_CODE_MAP = {
    'english': 'en',
    'spanish': 'es',
    'french': 'fr',
    'german': 'de',
    'portuguese': 'pt',
    'russian': 'ru',
    'italian': 'it',
}


class DocumentTranslator:
    """
    Document Translation Agent using Ollama.

    A class for translating documents (PDF, DOCX, TXT) between languages
    using Ollama's translation models. Supports chunking large documents
    and maintains sentence boundaries during translation.

    Attributes:
        context_ratio (float): Ratio of model context length to use (default 0.7).
        model (str): Ollama model to use for translation (default 'translate-gemma-27b').
        supported_formats (list): List of supported file formats.
        language_code_map (dict): Mapping of language names to codes.
    """

    def __init__(self, context_ratio: float = 0.7, model: str = 'translate-gemma-27b'):
        """
        Initialize the DocumentTranslator.

        Args:
            context_ratio (float): Ratio of model context length to use.
                                   Must be between 0.0 and 1.0. Default is 0.7.
            model (str): Ollama model to use for translation.
                         Default is 'translate-gemma-27b'.

        Raises:
            TranslationException: If context_ratio is not between 0.0 and 1.0.
        """
        method = "__init__"
        logging.debug(f"{method} - START - context_ratio={context_ratio}, model={model}")

        if not 0.0 < context_ratio <= 1.0:
            raise TranslationException(
                f"context_ratio must be between 0.0 and 1.0, got {context_ratio}"
            )

        self.context_ratio = context_ratio
        self.model = model
        self.supported_formats = SUPPORTED_FORMATS.copy()
        self.language_code_map = LANGUAGE_CODE_MAP.copy()

        logging.debug(f"{method} - END")

    def set_context_ratio(self, ratio: float):
        """
        Set the context ratio for translation.

        Args:
            ratio (float): New context ratio value. Must be between 0.0 and 1.0.

        Raises:
            TranslationException: If ratio is not between 0.0 and 1.0.
        """
        method = "set_context_ratio"
        logging.debug(f"{method} - Setting context_ratio to {ratio}")

        if not 0.0 < ratio <= 1.0:
            raise TranslationException(
                f"context_ratio must be between 0.0 and 1.0, got {ratio}"
            )

        self.context_ratio = ratio
        logging.debug(f"{method} - END")

    def get_supported_formats(self) -> List[str]:
        """
        Get the list of supported file formats.

        Returns:
            list: List of supported file format extensions (e.g., ['pdf', 'docx', 'txt']).
        """
        method = "get_supported_formats"
        logging.debug(f"{method} - START")
        logging.debug(f"{method} - Returning: {self.supported_formats}")
        return self.supported_formats.copy()

    def get_supported_languages(self) -> Dict[str, str]:
        """
        Get the mapping of supported language names to codes.

        Returns:
            dict: Dictionary mapping language names to their standard codes.
        """
        method = "get_supported_languages"
        logging.debug(f"{method} - START")
        logging.debug(f"{method} - Returning: {self.language_code_map}")
        return self.language_code_map.copy()

    def get_language_code(self, language_name: str) -> str:
        """
        Get the standard language code for a language name.

        Args:
            language_name (str): Language name (e.g., 'Spanish', 'English').

        Returns:
            str: Standard language code (e.g., 'es', 'en').

        Raises:
            TranslationException: If the language name is not supported.
        """
        method = "get_language_code"
        logging.debug(f"{method} - Looking up language: {language_name}")

        # Try direct lookup (case-insensitive)
        language_lower = language_name.lower().strip()

        if language_lower in self.language_code_map:
            code = self.language_code_map[language_lower]
            logging.debug(f"{method} - Found code: {code}")
            return code

        # If it's already a valid code, return it
        if len(language_name) == 2 and language_name.isalpha():
            logging.debug(f"{method} - Assuming '{language_name}' is a valid code")
            return language_name.lower()

        raise TranslationException(
            f"Unsupported language: '{language_name}'. "
            f"Supported languages: {', '.join(self.language_code_map.keys())}"
        )

    def _read_file(self, input_path: str) -> str:
        """
        Read and extract text from a file based on its format.

        Args:
            input_path (str): Path to the input file.

        Returns:
            str: Extracted text content from the file.

        Raises:
            TranslationUnsupportedFormatError: If the file format is not supported.
            TranslationException: If there's an error reading the file.
        """
        method = "_read_file"
        logging.debug(f"{method} - Reading file: {input_path}")

        file_path = Path(input_path)
        extension = file_path.suffix.lower().lstrip('.')

        if extension not in self.supported_formats:
            raise TranslationUnsupportedFormatError(format=extension)

        try:
            if extension == 'txt':
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                logging.debug(f"{method} - TXT file read successfully")
                return text

            elif extension == 'pdf':
                return self._read_pdf(input_path)

            elif extension == 'docx':
                return self._read_docx(input_path)

        except FileNotFoundError:
            raise TranslationException(f"File not found: {input_path}")
        except Exception as e:
            raise TranslationException(f"Error reading file: {str(e)}")

    def _read_pdf(self, input_path: str) -> str:
        """
        Read text from a PDF file.

        Args:
            input_path (str): Path to the PDF file.

        Returns:
            str: Extracted text content.
        """
        method = "_read_pdf"
        logging.debug(f"{method} - Reading PDF: {input_path}")

        try:
            import PyPDF2
            text = ""
            with open(input_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
            logging.debug(f"{method} - PDF read successfully")
            return text
        except ImportError:
            raise TranslationException(
                "PyPDF2 is not installed. Install it with: pip install PyPDF2"
            )

    def _read_docx(self, input_path: str) -> str:
        """
        Read text from a DOCX file.

        Args:
            input_path (str): Path to the DOCX file.

        Returns:
            str: Extracted text content.
        """
        method = "_read_docx"
        logging.debug(f"{method} - Reading DOCX: {input_path}")

        try:
            import docx
            doc = docx.Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            logging.debug(f"{method} - DOCX read successfully")
            return text
        except ImportError:
            raise TranslationException(
                "python-docx is not installed. Install it with: pip install python-docx"
            )

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks at sentence boundaries.

        Args:
            text (str): Text to chunk.

        Returns:
            list: List of text chunks, each within the context ratio limit.
        """
        method = "_chunk_text"
        logging.debug(f"{method} - Chunking text")

        # Estimate max chunk size (assuming ~4 chars per token, 8192 context)
        max_chunk_size = int(self.context_ratio * 8192 * 4)

        # Split at sentence boundaries (periods followed by space or end)
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)

        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if not sentence.strip():
                continue

            if len(current_chunk) + len(sentence) < max_chunk_size:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence

        if current_chunk:
            chunks.append(current_chunk)

        logging.debug(f"{method} - Created {len(chunks)} chunks")
        return chunks

    def _translate_chunk(self, chunk: str, target_lang: str) -> str:
        """
        Translate a single text chunk using Ollama.

        Args:
            chunk (str): Text chunk to translate.
            target_lang (str): Target language code.

        Returns:
            str: Translated text.

        Raises:
            TranslationAPIError: If there's an error with the Ollama API.
        """
        method = "_translate_chunk"
        logging.debug(f"{method} - Translating chunk to {target_lang}")

        prompt = f"Translate the following text to {target_lang}. Only return the translated text, nothing else:\n\n{chunk}"

        try:
            response = call_ollama_api_request(
                model=self.model,
                message=prompt,
                tipo='generate',
                temperature=0.3
            )

            if response.status_code == 200:
                translated_text = procesa_output(response)
                logging.debug(f"{method} - Chunk translated successfully")
                return translated_text.strip()
            else:
                raise TranslationAPIError(
                    status_code=response.status_code
                )

        except OllamaException as e:
            raise TranslationAPIError(message=str(e))
        except Exception as e:
            raise TranslationAPIError(message=str(e))

    def _save_text(self, text: str, output_path: str):
        """
        Save translated text to output file.

        Args:
            text (str): Translated text content.
            output_path (str): Path to save the translated file.
        """
        method = "_save_text"
        logging.debug(f"{method} - Saving to: {output_path}")

        output_file = Path(output_path)
        extension = output_file.suffix.lower().lstrip('.')

        if extension == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif extension == 'pdf':
            # For PDF, we'd need a PDF library to create the file
            # For now, save as TXT with a note
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logging.warning(f"{method} - PDF output saved as TXT format")
        elif extension == 'docx':
            try:
                import docx
                doc = docx.Document()
                doc.add_paragraph(text)
                doc.save(output_path)
            except ImportError:
                # Fallback to TXT
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                logging.warning(f"{method} - DOCX output saved as TXT format")
        else:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)

    def translate_document(
        self,
        input_path: str,
        output_path: str,
        target_lang: str
    ):
        """
        Translate a document from source language to target language.

        This is the main method for document translation. It reads the
        input file, chunks it if necessary, translates each chunk using
        Ollama, and saves the translated content to the output file.

        Args:
            input_path (str): Path to the source document (PDF, DOCX, or TXT).
            output_path (str): Path to save the translated document.
            target_lang (str): Target language code (e.g., 'es') or name
                             (e.g., 'Spanish').

        Raises:
            TranslationUnsupportedFormatError: If the input file format is not supported.
            TranslationException: If there's an error during translation.
            TranslationAPIError: If there's an error with the Ollama API.

        Example:
            >>> translator = DocumentTranslator()
            >>> translator.translate_document(
            ...     "input_document.pdf",
            ...     "translated_es.pdf",
            ...     "Spanish"
            ... )
        """
        method = "translate_document"
        logging.debug(f"{method} - START - input={input_path}, output={output_path}, lang={target_lang}")

        # Get target language code
        target_code = self.get_language_code(target_lang)
        logging.debug(f"{method} - Target language code: {target_code}")

        # Read the input file
        text = self._read_file(input_path)
        logging.debug(f"{method} - File read, length: {len(text)} characters")

        # Check if chunking is needed
        max_chunk_size = int(self.context_ratio * 8192 * 4)
        if len(text) > max_chunk_size:
            logging.debug(f"{method} - Document requires chunking")
            chunks = self._chunk_text(text)
            translated_chunks = []

            for i, chunk in enumerate(chunks):
                logging.debug(f"{method} - Translating chunk {i + 1}/{len(chunks)}")
                translated_chunk = self._translate_chunk(chunk, target_code)
                translated_chunks.append(translated_chunk)

            translated_text = "\n\n".join(translated_chunks)
        else:
            logging.debug(f"{method} - Translating full document")
            translated_text = self._translate_chunk(text, target_code)

        # Save the translated document
        self._save_text(translated_text, output_path)
        logging.debug(f"{method} - END - Document translated successfully")
